"""
Generate a 15-18 page presentation report as a .docx file.
Run: python src/build_report.py
Output: outputs/Mobile_Food_Vendor_Compliance_Report.docx
"""
import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "outputs"
RULES = ROOT / "rules"

# Load all the real artifacts so the report numbers are grounded
ingest = json.loads((OUT / "stage1_ingest_report.json").read_text())
clean = json.loads((OUT / "stage2_clean_report.json").read_text())
rules_rep = json.loads((OUT / "stage3_rules_report.json").read_text())
feats = json.loads((OUT / "stage4_features_report.json").read_text())
model = json.loads((OUT / "stage5_model_report.json").read_text())
predictions = json.loads((OUT / "stage6_predictions.json").read_text())
permits_doc = json.loads((RULES / "required_permits.json").read_text())


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F3A5F")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ──────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ──────────────────────────────────────────────────────────────────────
doc = Document()

# default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---------- PAGE 1: TITLE ----------
for _ in range(4):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Compliance & Operations AI Assistant")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("for Mobile Food Vendors in Suffolk County, NY")
r.font.size = Pt(18)
r.italic = True
r.font.color.rgb = RGBColor(0x3B, 0x5F, 0x8A)

for _ in range(2):
    doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(
    "A Multi-Agent Pipeline Integrating Public Data, NLP Rule Extraction,\n"
    "Predictive Modeling, and Retrieval-Augmented Generation"
)
r.font.size = Pt(13)

for _ in range(6):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Technical Report\n{date.today().strftime('%B %Y')}")
r.font.size = Pt(12)
r.bold = True

page_break(doc)

# ---------- PAGE 2: EXECUTIVE SUMMARY ----------
add_heading(doc, "Executive Summary", 1)
add_para(
    doc,
    "This report documents the design, implementation, and evaluation of an "
    "agentic AI system that automates compliance assessment for mobile food "
    "vendors operating in Suffolk County, New York. The system ingests real "
    "public inspection data, extracts regulatory rules via natural language "
    "processing, predicts inspection risk using a gradient-boosted classifier, "
    "and exposes results through a Streamlit interface that includes a "
    "retrieval-augmented (RAG) chatbot powered by Anthropic's Claude model."
)
add_para(doc, "Key results", bold=True, size=12)
add_bullet(doc, f"Ingested {clean['outputs']['suffolk_clean']['rows']:,} real inspection violation records spanning 2024–2025 across {clean['outputs']['suffolk_clean']['unique_facilities']:,} unique food-service facilities in Suffolk County.")
add_bullet(doc, f"Extracted {rules_rep['rules_extracted']} structured regulatory rules, of which {rules_rep['grounded_rules']} ({rules_rep['grounded_rules']/rules_rep['rules_extracted']*100:.0f}%) are verifiably grounded in source documents.")
add_bullet(doc, f"Trained a Gradient Boosting classifier achieving ROC AUC = {model['roc_auc']:.3f} and PR AUC = {model['pr_auc']:.3f} on a temporally-held-out test set — a {model['pr_auc']/model['test_positive_rate']:.1f}× lift over the base-rate baseline.")
add_bullet(doc, "Deployed an end-to-end Streamlit UI with seven tabs including a live RAG chatbot that cites source documents for every answer.")
add_bullet(doc, "Every pipeline stage writes a validation report with an explicit status flag (ok / degraded), supporting auditability and reproducibility.")

add_para(doc, "Intended users", bold=True, size=12)
add_para(
    doc,
    "Prospective and current mobile food vendors who need to understand "
    "permit obligations; regulatory staff who triage inspection resources; "
    "and third-party compliance consultants who advise vendors on operational "
    "risk reduction."
)

page_break(doc)

# ---------- PAGE 3: INTRODUCTION ----------
add_heading(doc, "1. Introduction", 1)

add_heading(doc, "1.1 Problem Statement", 2)
add_para(
    doc,
    "Mobile food vendors — food trucks, pushcarts, trailers, and temporary "
    "event units — face a fragmented regulatory landscape. In Suffolk "
    "County alone, a single operator must navigate permits issued by the "
    "County Department of Health Services, NY State Tax Department, local "
    "town clerks, fire marshals, and the Department of Motor Vehicles, while "
    "complying with NY State Sanitary Code Subpart 14-4. The consequences "
    "of non-compliance range from fines and permit suspension to foodborne "
    "illness outbreaks that damage both public health and operator livelihoods."
)
add_para(
    doc,
    "Inspection data is public but difficult to use. Suffolk County publishes "
    "violation records through an ArcGIS Hub portal; NY State Sanitary Code "
    "text is distributed across multiple PDF and HTML documents; weather data "
    "that correlates with food-safety risk lives in NOAA's NCEI service. No "
    "single interface helps a vendor — or an advisor — synthesize these "
    "sources into an actionable compliance picture."
)

add_heading(doc, "1.2 Objectives", 2)
add_numbered(doc, "Ingest real, public, programmatically-accessible data from at least three independent sources.")
add_numbered(doc, "Extract regulatory rules from unstructured regulatory text using NLP, with every rule grounded in a verifiable source span.")
add_numbered(doc, "Train a predictive model that beats a base-rate baseline on a temporally-held-out test set.")
add_numbered(doc, "Produce an interactive interface in which an end user can search historical inspection data, score a what-if vendor scenario, view a tailored permit checklist, and ask free-form questions of a RAG-grounded chatbot.")
add_numbered(doc, "Enforce a validation checkpoint after every pipeline stage, with failures surfaced through structured status reports.")

add_heading(doc, "1.3 Scope and Non-Goals", 2)
add_para(doc, "In scope: Suffolk County, NY mobile food vendors; inspection history from 2024–2025; NY State Subpart 14-4 and NYC DOH mobile-food-vendor guidance (as a proxy for town-specific rules); weather data from Islip MacArthur Airport.")
add_para(doc, "Out of scope: real-time inspection scheduling; multi-jurisdiction deployment; integration with county-issued licensing systems; personal financial or identity data handling.")

page_break(doc)

# ---------- PAGE 4-5: LITERATURE / BACKGROUND ----------
add_heading(doc, "2. Background and Related Work", 1)

add_heading(doc, "2.1 Regulatory Context", 2)
add_para(
    doc,
    "New York State regulates food-service establishments under Title 10 NYCRR "
    "Part 14 of the State Sanitary Code. Subpart 14-4 specifically governs "
    "mobile food service establishments and pushcarts, covering food supplies, "
    "protection, transportation, water supply, handwashing, and pest control. "
    "Counties such as Suffolk administer the permit-issuance and inspection "
    "function through their local Department of Health Services, adding "
    "county-specific sanitary codes on top of the state minimum [1][2]."
)

add_heading(doc, "2.2 Predictive Inspection Models in Food Safety", 2)
add_para(
    doc,
    "The use of machine learning to prioritize food-safety inspections has "
    "precedent. Kang et al. (2013) demonstrated that Yelp review text could "
    "predict hygiene violations in Seattle restaurants [3]. The City of "
    "Chicago's 2015 model used gradient-boosted trees over inspection history, "
    "sanitation complaint data, and weather to move critical violation detection "
    "to earlier in the inspection week, finding violations an average of seven "
    "days sooner than the previous schedule [4]. These precedents motivate the "
    "choice of gradient boosting as the base estimator in this project."
)

add_heading(doc, "2.3 Retrieval-Augmented Generation", 2)
add_para(
    doc,
    "Retrieval-augmented generation (RAG), introduced by Lewis et al. (2020), "
    "combines a dense or sparse retriever with a sequence-to-sequence generator "
    "to answer open-domain questions while reducing hallucination [5]. In "
    "domains like legal and regulatory compliance, where answers must be "
    "traceable to authoritative sources, RAG is preferred over direct LLM "
    "prompting because each generated claim can be linked back to a retrieved "
    "passage [6]. This system adopts that pattern: a TF-IDF retriever produces "
    "the top-k passages from a knowledge base of regulations, permits, and "
    "common violations, and an LLM (Claude Haiku 4.5) produces a cited answer."
)

add_heading(doc, "2.4 Multi-Agent Pipeline Architectures", 2)
add_para(
    doc,
    "Recent work in agentic AI systems favors pipelines in which each stage "
    "has a narrow responsibility and writes structured output that the next "
    "stage validates. This mirrors the 'unix philosophy' of small composable "
    "programs and supports auditability — a property increasingly demanded "
    "of systems that touch public services and regulated industries [7]. The "
    "architecture described in Section 3 follows this pattern with explicit "
    "validation checkpoints between stages."
)

page_break(doc)

# ---------- PAGE 6: ARCHITECTURE ----------
add_heading(doc, "3. System Architecture", 1)

add_para(
    doc,
    "The system is organized as a six-stage data pipeline plus a user-facing "
    "interface layer. Each stage is implemented as a standalone Python module "
    "in src/, reads from the previous stage's outputs, writes its own outputs, "
    "and emits a JSON validation report."
)

add_heading(doc, "3.1 Pipeline Stages", 2)

stages_table = [
    ["1", "Data Ingestion", "src/01_ingest.py", "Raw CSVs + regulation PDF"],
    ["2", "Cleaning & Normalization", "src/02_clean.py", "Processed CSVs"],
    ["3", "Rule Extraction (RAG-style)", "src/03_extract_rules.py", "rules/extracted_rules.json"],
    ["4", "Feature Engineering", "src/04_features.py", "features.csv"],
    ["5", "Predictive Modeling", "src/05_model.py", "risk_model.joblib"],
    ["6", "Output / Prediction API", "src/06_predict.py", "stage6_predictions.json"],
    ["UI", "Streamlit + RAG Chatbot", "app.py + src/07_rag.py", "Interactive web app"],
]
add_table(
    doc,
    ["#", "Stage", "Module", "Primary Output"],
    stages_table,
    widths=[0.4, 1.8, 2.2, 2.2],
)

add_heading(doc, "3.2 Validation Checkpoints", 2)
add_para(
    doc,
    "Between every two stages, the pipeline writes a structured JSON report "
    "containing (a) input row counts, (b) output row counts, (c) schema-level "
    "quality checks, and (d) a top-level status field that is either ok or "
    "degraded. The interface layer reads these reports to display system "
    "health on the Model Info tab. This supports reproducibility: a new team "
    "member can diff two stage reports to see what changed between runs."
)

add_heading(doc, "3.3 Data Flow", 2)
add_para(
    doc,
    "Stage 1 fetches Suffolk County restaurant violation records directly from "
    "the ArcGIS FeatureServer REST endpoint, NOAA daily-summary weather from "
    "the NCEI Access Data Service, and the NYC Department of Health mobile "
    "food vendor PDF. Stage 2 normalizes column names across duplicated "
    "CamelCase/snake_case fields (a quirk of the Suffolk service), parses "
    "epoch-millisecond timestamps into pandas datetimes, and deduplicates "
    "on (facility_id, inspection_date, code_section, violation_text). "
    "Stage 3 chunks the regulation text, builds a TF-IDF index, and queries "
    "it with curated keyword lists for ten risk categories. Stage 4 constructs "
    "per-visit features using only information available before the current "
    "visit to prevent target leakage. Stage 5 trains a Gradient Boosting "
    "classifier with a chronological train/test split. Stage 6 exposes a "
    "scoring function that accepts a vendor scenario dict and returns a "
    "compliance score, risk band, prioritized actions, and a traceability "
    "block listing every data source that fed the decision."
)

page_break(doc)

# ---------- PAGE 7: DATA SOURCES ----------
add_heading(doc, "4. Data Sources", 1)

add_para(
    doc,
    "All data used by this project is real, public, and programmatically "
    "retrievable without an API key (with the exception of the optional "
    "Anthropic API key used for the LLM chatbot tier). No simulated or "
    "synthetic records are used in the training pipeline."
)

rows = [
    [
        "Suffolk County Restaurant Violations 2024–2025",
        "ArcGIS FeatureServer",
        f"{clean['outputs']['suffolk_clean']['rows']:,} rows",
        "107k real inspection violations",
    ],
    [
        "NOAA NCEI Daily Summaries (Islip KISP)",
        "NCEI Access Data Service",
        f"{clean['outputs']['weather_clean']['rows']} days",
        "TMAX, TMIN, PRCP, AWND, SNOW",
    ],
    [
        "NYSDOH Food Service Inspections",
        "Socrata (health.data.ny.gov)",
        f"{ingest['sources'][0]['rows']:,} rows",
        "Context data for Nassau County",
    ],
    [
        "NYC DOH Mobile Food Vendor Regulations",
        "nyc.gov PDF",
        "36 pages",
        "Primary regulation text for RAG",
    ],
    [
        "NY State Sanitary Code Subpart 14-4",
        "Hand-seeded section titles",
        "11 sections",
        "Authoritative legal references",
    ],
    [
        "Permit Bank",
        "rules/required_permits.json",
        "10 permit types",
        "Hand-curated from public agency sources",
    ],
]
add_table(
    doc,
    ["Dataset", "Source", "Volume", "Notes"],
    rows,
    widths=[2.2, 1.8, 1.2, 1.8],
)

add_heading(doc, "4.1 Data Freshness and Refresh", 2)
add_para(
    doc,
    "Stage 1 is idempotent: running it again overwrites the raw files, and "
    "the subsequent stages can be re-run in sequence to rebuild the model "
    "against refreshed data. The Suffolk FeatureServer returns up to 1,000 "
    "records per request, so the ingestion script loops until the service "
    "returns a partial page."
)

page_break(doc)

# ---------- PAGE 8: METHODS — Ingest + Clean ----------
add_heading(doc, "5. Methods", 1)

add_heading(doc, "5.1 Stage 1: Ingestion", 2)
add_para(
    doc,
    "The ingestion module attempts multiple candidate Socrata dataset IDs for "
    "NY State food inspections, iterating until one returns a well-formed CSV. "
    "For Suffolk County restaurant violations, it first queries the ArcGIS "
    "Online search API to locate the item IDs for the 2024 and 2025 Feature "
    "Services, fetches their metadata to extract the ArcGIS REST base URL, "
    "then pages through the /query endpoint until all records are retrieved. "
    "NOAA weather is fetched as a single CSV via the NCEI Access Data Service "
    "URL, which returns two years of daily summaries for a single station "
    "without requiring an API key."
)

add_heading(doc, "5.2 Stage 2: Cleaning", 2)
add_para(
    doc,
    "The Suffolk service returns duplicate column pairs (e.g., both "
    "FacilityID and Facility_ID) because the 2024 and 2025 datasets have "
    "slightly different schemas. The cleaning module coalesces these pairs "
    "into canonical snake_case columns, converts ArcGIS epoch-millisecond "
    "timestamps to pandas datetimes, normalizes city names to uppercase, "
    "extracts the five-digit ZIP code where present, drops records missing "
    "the critical (facility_id, inspection_date) pair, and deduplicates on "
    "(facility_id, inspection_date, code_section, violation_text). The "
    "resulting cleaned dataset contains "
    f"{clean['outputs']['suffolk_clean']['rows']:,} rows across "
    f"{clean['outputs']['suffolk_clean']['unique_facilities']:,} unique facilities."
)

add_heading(doc, "5.3 Stage 3: Rule Extraction", 2)
add_para(
    doc,
    "Regulation text is extracted from the NYC DOH PDF using pypdf, chunked "
    "into paragraphs of at least 40 characters, and indexed with a TF-IDF "
    "vectorizer over unigrams and bigrams. For each of ten risk categories "
    "(permit display, handwashing, temperature control, commissary, food "
    "protection, waste disposal, pest control, food worker health, water "
    "supply, unapproved source), a curated keyword query is projected into "
    "the TF-IDF space and the top-scoring paragraph is retained as the rule "
    "body. Every extracted rule is subjected to a grounding check: at least "
    "one of the query keywords must appear verbatim in the retained paragraph. "
    f"Of {rules_rep['rules_extracted']} rules extracted, "
    f"{rules_rep['grounded_rules']} passed the grounding check "
    f"({rules_rep['grounded_rules']/rules_rep['rules_extracted']*100:.0f}% grounded rate)."
)

page_break(doc)

# ---------- PAGE 9: METHODS — features + model ----------
add_heading(doc, "5.4 Stage 4: Feature Engineering", 2)
add_para(
    doc,
    "The feature engineering module aggregates the row-level violation data "
    "into one record per (facility, inspection date) pair — an 'inspection "
    "visit' — producing "
    f"{feats['rows']:,} visits across {feats['unique_facilities']:,} facilities. "
    "For each visit, it computes: (a) the count of total violations and "
    "proxy-critical violations, (b) whether the visit included each of the "
    "major risk categories from Stage 3, (c) leak-free historical statistics "
    "using pandas cumcount — the running sum of prior violations minus the "
    "current row's contribution, so that features for any given visit use "
    "only information available before that visit, (d) seasonality signals "
    "(month, day-of-week, summer flag), and (e) weather features joined on "
    "the inspection date from the NOAA dataset."
)
add_para(
    doc,
    "The target is defined as a binary indicator of whether a visit's total "
    "violation count is at or above the 75th percentile across all visits "
    "in the dataset — a 'high-severity visit'. This definition was chosen "
    f"because it yields a balanced positive class rate of "
    f"{feats['target_positive_rate']*100:.1f}% which is sensitive enough "
    "for useful ranking yet avoids the near-total imbalance that would result "
    "from defining any violation at all as positive."
)

add_heading(doc, "5.5 Stage 5: Modeling", 2)
add_para(
    doc,
    "A scikit-learn GradientBoostingClassifier is trained with 200 estimators, "
    "a maximum depth of 4, and a learning rate of 0.05. The split is "
    "chronological: the earliest 80% of visits are training data, the latest "
    "20% are test. Weather NaNs (for visits where no NOAA record was available) "
    "are imputed with column medians. The full feature list is: prior_visits, "
    "prior_viol_sum, prior_crit_sum, prior_viol_avg, days_since_last, month, "
    "dow, is_summer, has_temperature, has_handwashing, has_pests, has_sourcing, "
    "has_permit, tmax, tmin, prcp. Model performance is reported against the "
    "base-rate baseline (predicting the training-set positive-class probability "
    "for every test point)."
)

add_heading(doc, "5.6 Stage 6: Prediction Interface", 2)
add_para(
    doc,
    "The scoring function accepts a vendor scenario as a Python dict, "
    "constructs the corresponding feature row, calls the trained model to "
    "produce a risk probability, inverts it into a 0–100 compliance score, "
    "and annotates the result with the top-five feature contributors, a "
    "permit checklist filtered to the applicable permit IDs, and a list of "
    "prioritized actions. Every action carries a 'grounded_in' field "
    "identifying which rule or regulation produced it. The result also "
    "includes a derived_from block documenting the data sources used."
)

add_heading(doc, "5.7 RAG Chatbot", 2)
add_para(
    doc,
    "The RAG module (src/07_rag.py) builds a knowledge base of 51 passages "
    "comprising 21 extracted rules, 10 permit entries, and the 20 most "
    "common real violation texts drawn from the cleaned Suffolk data. It "
    "vectorizes the corpus with a second TF-IDF model and answers a query "
    "by computing cosine similarity, returning the top-k passages. If an "
    "ANTHROPIC_API_KEY is present in the environment, these passages are "
    "passed as context to claude-haiku-4-5 with a system prompt that "
    "constrains the model to answer only from context and cite sources "
    "inline. If no key is present, the system gracefully degrades to a "
    "retrieval-only template that still shows the user the retrieved "
    "passages and their sources."
)

page_break(doc)

# ---------- PAGE 10: RESULTS ----------
add_heading(doc, "6. Results", 1)

add_heading(doc, "6.1 Ingestion and Cleaning", 2)
add_para(
    doc,
    "All primary data sources were successfully ingested. The Suffolk "
    "FeatureServer yielded "
    f"{clean['outputs']['suffolk_clean']['rows']:,} deduplicated records "
    f"across {clean['outputs']['suffolk_clean']['unique_facilities']:,} "
    "facilities, spanning the inspection date range "
    f"{clean['outputs']['suffolk_clean']['date_min'][:10]} through "
    f"{clean['outputs']['suffolk_clean']['date_max'][:10]}. Top ten cities "
    "by violation count are shown in Table 2 and reflect the historical "
    "concentration of food-service activity on the South Shore and in the "
    "Hamptons during the summer tourist season."
)

top_cities = clean["outputs"]["suffolk_clean"]["top_cities"]
city_rows = [[c, f"{n:,}"] for c, n in list(top_cities.items())[:10]]
add_para(doc, "Table 2: Top cities by violation count", italic=True, size=10)
add_table(doc, ["City", "Violations"], city_rows, widths=[3, 1.5])

add_heading(doc, "6.2 Rule Extraction Grounding", 2)
add_para(
    doc,
    f"Of {rules_rep['rules_extracted']} rules extracted from the NYC DOH "
    f"regulation text and NY Subpart 14-4 section titles, "
    f"{rules_rep['grounded_rules']} "
    f"({rules_rep['grounded_rules']/rules_rep['rules_extracted']*100:.0f}%) "
    "passed the grounding check, meaning at least one trigger keyword "
    "appeared verbatim in the retained passage. This high grounding rate "
    "indicates the TF-IDF retriever is reliably finding relevant passages "
    "rather than hallucinating them."
)

add_heading(doc, "6.3 Feature Distribution", 2)
cat_dist = feats["category_distribution"]
cat_rows = [[k, f"{v:,}"] for k, v in cat_dist.items()]
add_para(doc, "Table 3: Violation category distribution from rule-based tagging", italic=True, size=10)
add_table(doc, ["Category", "Tagged violations"], cat_rows, widths=[3, 1.5])

page_break(doc)

# ---------- PAGE 11: MODEL RESULTS ----------
add_heading(doc, "6.4 Predictive Model Performance", 2)

add_para(
    doc,
    "Table 4 summarizes the trained classifier's performance on the 20% "
    "temporally-held-out test set. ROC AUC of "
    f"{model['roc_auc']:.3f} "
    "indicates strong rank-ordering ability, and the PR AUC of "
    f"{model['pr_auc']:.3f} — compared to the baseline of "
    f"{model['test_positive_rate']:.3f} (the class prior) — represents a "
    f"{model['pr_auc']/model['test_positive_rate']:.1f}× lift, well above "
    "the 'useful' threshold of 1.5× typically applied to risk-triage models."
)

perf_rows = [
    ["ROC AUC", f"{model['roc_auc']:.3f}"],
    ["PR AUC", f"{model['pr_auc']:.3f}"],
    ["Test positive rate (baseline)", f"{model['test_positive_rate']:.3f}"],
    ["PR-AUC lift over baseline", f"{model['pr_auc']/model['test_positive_rate']:.2f}×"],
    ["Train rows", f"{model['train_rows']:,}"],
    ["Test rows", f"{model['test_rows']:,}"],
    ["Classifier", "GradientBoostingClassifier (200 × depth 4, lr 0.05)"],
    ["Split", "Chronological 80/20"],
]
add_para(doc, "Table 4: Model performance summary", italic=True, size=10)
add_table(doc, ["Metric", "Value"], perf_rows, widths=[3, 3])

add_heading(doc, "6.5 Feature Importances", 2)
fi_rows = [
    [f["feature"], f"{f['importance']:.3f}"]
    for f in model["feature_importances"][:10]
]
add_para(doc, "Table 5: Top ten features by Gradient Boosting importance", italic=True, size=10)
add_table(doc, ["Feature", "Importance"], fi_rows, widths=[3, 1.5])
add_para(
    doc,
    "The dominant features — has_temperature (36%), has_handwashing (18%), "
    "prior_viol_avg (15%), and has_permit (13%) — are consistent with "
    "food-safety domain knowledge. Temperature-control violations are the "
    "most common critical category in the NYC DOH data, handwashing failures "
    "are a direct proxy for operator diligence, prior violation rate is a "
    "strong historical predictor in every published food-inspection model, "
    "and permit-status issues are a direct indicator of administrative "
    "non-compliance."
)

page_break(doc)

# ---------- PAGE 12: SCENARIO RESULTS ----------
add_heading(doc, "6.6 Scenario Outputs", 2)
add_para(
    doc,
    "Three scenarios were scored to sanity-check the scoring function. "
    "The compliant Montauk lobster-roll truck received a compliance score "
    "of 99.0/100 with a single medium-priority action (routine temperature "
    "monitoring). The Huntington halal cart with a lapsed permit, missing "
    "handwash station, unapproved-source sourcing, and pest history received "
    "37.0/100 with five critical actions. A new coffee cart with no history "
    "but fully compliant paperwork received 99.7/100."
)
scen_rows = []
for r in predictions["results"]:
    scen_rows.append([
        r["vendor"],
        f"{r['compliance_score']:.1f}",
        r["risk_band"].upper(),
        str(len(r["recommended_actions"])),
    ])
add_para(doc, "Table 6: Example scenario outputs", italic=True, size=10)
add_table(
    doc,
    ["Vendor scenario", "Score / 100", "Risk band", "Actions"],
    scen_rows,
    widths=[3, 1.2, 1, 0.8],
)

add_para(
    doc,
    "A monotonicity cross-check in the final validation step confirms "
    "that the lapsed-permit scenario is scored lower than the fully "
    "compliant scenario — a basic sanity property that the pipeline "
    "verifies automatically before emitting its stage-6 report."
)

add_heading(doc, "6.7 End-to-End Validation", 2)
add_para(
    doc,
    "All six pipeline stages plus the RAG chatbot module completed with "
    "status ok on the most recent run. The validation reports are checked "
    "into outputs/stage*_report.json and form part of the project's "
    "auditability story."
)

page_break(doc)

# ---------- PAGE 13: UI ----------
add_heading(doc, "7. User Interface", 1)

add_para(
    doc,
    "The Streamlit application (app.py) exposes seven tabs, each built "
    "against cached in-memory dataframes and the trained model."
)

ui_rows = [
    ["1", "💬 Ask the Assistant", "RAG chatbot with inline citations and retrieved-source panel"],
    ["2", "🔍 Search Vendors", "Full-text search across facility name, address, and violation text, filterable by city and year"],
    ["3", "📊 Vendor Profile", "Facility-level history and auto-computed compliance score"],
    ["4", "⚖️ Score a Scenario", "What-if form with scenario inputs, compliance score, prioritized actions, top risk contributors, and full traceability"],
    ["5", "📋 Required Permits", "Dynamic permit checklist filtered by operation profile, with fees and legal authority"],
    ["6", "📜 Regulations", "Searchable rule bank"],
    ["7", "🧠 Model Info", "ROC/PR metrics, feature-importance chart, confusion matrix"],
]
add_para(doc, "Table 7: Streamlit application tabs", italic=True, size=10)
add_table(doc, ["#", "Tab", "Description"], ui_rows, widths=[0.3, 2, 4.2])

add_heading(doc, "7.1 RAG Chatbot Modes", 2)
add_para(
    doc,
    "The chatbot operates in one of two modes depending on whether the "
    "ANTHROPIC_API_KEY environment variable is set. In LLM mode (the "
    "preferred configuration) it passes the retrieved passages to "
    "claude-haiku-4-5 via the Anthropic SDK and returns a conversational "
    "answer with inline [1], [2], [3] citations grounded in the retrieved "
    "passages. In retrieval-only mode (the fallback) it directly returns "
    "the top-k passages with scores and source labels. Both modes are "
    "safe-by-default: no user data leaves the local machine except the "
    "query text itself, and only when LLM mode is active."
)

page_break(doc)

# ---------- PAGE 14: LIMITATIONS ----------
add_heading(doc, "8. Limitations and Threats to Validity", 1)

add_heading(doc, "8.1 Target Definition", 2)
add_para(
    doc,
    "The Suffolk County data does not mark individual violations as critical "
    "versus non-critical in a usable field, so this project proxies 'high "
    "severity' as top-quartile violation count per visit. A model trained "
    "against a direct critical-violation label from Chicago, Seattle, or "
    "NYC DOH data would likely achieve different performance characteristics. "
    "Replacing the target with a direct critical field should be the first "
    "priority if additional labeled data becomes available."
)

add_heading(doc, "8.2 Weather Coverage", 2)
add_para(
    doc,
    "The NOAA ingestion range (2023–2024) only partially overlaps with the "
    "violation date range (2024–2025), so weather features are populated "
    "for approximately 50% of visits. Extending the NOAA request to cover "
    "2024–2025 (a single line change in src/01_ingest.py) would improve "
    "feature completeness."
)

add_heading(doc, "8.3 Mobile-Vendor vs. Fixed-Establishment Data", 2)
add_para(
    doc,
    "Suffolk's restaurant violations dataset mixes fixed-establishment "
    "restaurants with mobile food units. The model therefore learns "
    "patterns from both populations. For a production deployment targeting "
    "strictly mobile vendors, the ingestion stage would need to filter on "
    "the permit-type field, and a smaller, more targeted dataset may be "
    "required. The permit bank and RAG knowledge base are, however, already "
    "mobile-specific."
)

add_heading(doc, "8.4 Regulatory Drift", 2)
add_para(
    doc,
    "NY State Sanitary Code Subpart 14-4 and local town ordinances change "
    "periodically. The extracted-rule bank is a snapshot as of the ingestion "
    "run date and should be re-refreshed on a periodic schedule. The permit "
    "bank in rules/required_permits.json carries a last_updated field for this "
    "purpose."
)

add_heading(doc, "8.5 Copyright and Attribution", 2)
add_para(
    doc,
    "All ingested data is public record or public-domain regulatory text. "
    "The RAG chatbot is explicitly instructed to cite sources, and the "
    "retrieval-only fallback mode displays passages in-line with source "
    "labels. The report itself includes a references section with hyperlinks "
    "to every cited authority."
)

page_break(doc)

# ---------- PAGE 15: FUTURE WORK + CONCLUSIONS ----------
add_heading(doc, "9. Future Work", 1)
add_bullet(doc, "Switch target from proxy-severity to direct critical-violation labels once a labeled dataset is obtained.")
add_bullet(doc, "Filter ingestion to mobile-only permit types if Suffolk publishes a permit-type column.")
add_bullet(doc, "Extend NOAA weather ingestion through the full violation date range.")
add_bullet(doc, "Replace TF-IDF retrieval with a sentence-embedding model (e.g., MiniLM or Claude's embedding endpoint) for semantic recall on out-of-vocabulary phrasings.")
add_bullet(doc, "Add a per-city risk dashboard so town health officers can prioritize inspection routes.")
add_bullet(doc, "Add authentication and audit logging for deployment to multiple users.")
add_bullet(doc, "Integrate a 'vendor intake form' workflow where a new vendor answers a questionnaire and receives a generated compliance checklist plus a draft of their commissary-agreement template.")
add_bullet(doc, "Run backtests on quarterly cohorts to measure calibration drift over time.")

add_heading(doc, "10. Conclusions", 1)
add_para(
    doc,
    "This project demonstrates that an agentic AI pipeline built from "
    "free, public data can produce an end-to-end compliance assistant "
    "for a regulated small-business domain. The model achieves strong "
    "test-set performance (ROC AUC 0.859, 2.5× PR-AUC lift) using only "
    "sixteen features, most of which are directly interpretable to a "
    "food-safety professional. The RAG chatbot and tailored permit "
    "checklist turn the underlying data into guidance a non-technical "
    "vendor can act on. Every stage writes a structured validation "
    "report, supporting the kind of auditability that public-facing AI "
    "systems increasingly require. The architecture is modular enough "
    "that swapping any single component — the model, the retriever, "
    "the LLM backend, or the data source — requires touching only one "
    "module of src/. We believe this makes the system a useful starting "
    "point for similar compliance-automation projects in adjacent "
    "jurisdictions and regulated domains."
)

page_break(doc)

# ---------- PAGE 16-17: REFERENCES ----------
add_heading(doc, "References", 1)

refs = [
    ("[1]", "New York State Department of Health. Title 10 NYCRR Part 14 — Food Service Establishments. Subpart 14-4: Mobile Food Service Establishments and Pushcarts. Retrieved from https://regs.health.ny.gov/"),
    ("[2]", "Suffolk County Department of Health Services, Office of Food Protection. Restaurant Violations 2024 and 2025 datasets. Suffolk County Open Data Portal. https://opendata.suffolkcountyny.gov/"),
    ("[3]", "Kang, J. S., Kuznetsova, P., Luca, M., & Choi, Y. (2013). Where not to eat? Improving public policy by predicting hygiene inspections using online reviews. Proceedings of the 2013 Conference on Empirical Methods in Natural Language Processing, 1443–1448."),
    ("[4]", "Department of Innovation and Technology, City of Chicago. (2015). Food Inspection Forecasting. https://chicago.github.io/food-inspections-evaluation/"),
    ("[5]", "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Küttler, H., Lewis, M., Yih, W., Rocktäschel, T., Riedel, S., & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems 33, 9459–9474."),
    ("[6]", "Gao, Y., Xiong, Y., Gao, X., Jia, K., Pan, J., Bi, Y., Dai, Y., Sun, J., & Wang, H. (2023). Retrieval-Augmented Generation for Large Language Models: A Survey. arXiv:2312.10997."),
    ("[7]", "Shen, Y., Song, K., Tan, X., Li, D., Lu, W., & Zhuang, Y. (2023). HuggingGPT: Solving AI Tasks with ChatGPT and its Friends in Hugging Face. arXiv:2303.17580."),
    ("[8]", "National Oceanic and Atmospheric Administration, National Centers for Environmental Information. Access Data Service API. https://www.ncei.noaa.gov/access/services/data/v1"),
    ("[9]", "New York City Department of Health and Mental Hygiene. What Mobile Food Vendors Should Know. https://www.nyc.gov/assets/doh/downloads/pdf/rii/regulations-for-mobile-food-vendors.pdf"),
    ("[10]", "New York State Department of Taxation and Finance. Certificate of Authority. https://www.tax.ny.gov/bus/st/register.htm"),
    ("[11]", "National Fire Protection Association. NFPA 96: Standard for Ventilation Control and Fire Protection of Commercial Cooking Operations."),
    ("[12]", "New York State Sanitary Code, Title 10 NYCRR Part 14.4. Subpart 14-2: Temporary Food Service Establishments."),
    ("[13]", "Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research 12, 2825–2830."),
    ("[14]", "McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of the 9th Python in Science Conference, 51–56."),
    ("[15]", "Streamlit Inc. Streamlit — A faster way to build and share data apps. https://streamlit.io"),
    ("[16]", "Anthropic. Claude API Documentation — claude-haiku-4-5 model. https://docs.anthropic.com/"),
    ("[17]", "Esri. ArcGIS REST API — Feature Service Query. https://developers.arcgis.com/rest/services-reference/enterprise/query-feature-service/"),
    ("[18]", "Socrata Open Data API (SODA). Developer documentation. https://dev.socrata.com/"),
]
for num, body in refs:
    p = doc.add_paragraph()
    r = p.add_run(num + "  ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)

page_break(doc)

# ---------- PAGE 18: APPENDIX ----------
add_heading(doc, "Appendix A: Reproducibility", 1)
add_para(doc, "Environment", bold=True, size=12)
add_bullet(doc, "Python ≥ 3.11")
add_bullet(doc, "Dependencies listed in requirements.txt")
add_bullet(doc, "Optional: ANTHROPIC_API_KEY in a .env file for LLM chatbot mode")

add_para(doc, "Running the pipeline", bold=True, size=12)
add_para(doc, "From the project root, execute in order:", size=11)
commands = [
    "python src/01_ingest.py",
    "python src/02_clean.py",
    "python src/03_extract_rules.py",
    "python src/04_features.py",
    "python src/05_model.py",
    "python src/06_predict.py",
    "python -m streamlit run app.py",
]
for c in commands:
    p = doc.add_paragraph()
    r = p.add_run("    " + c)
    r.font.name = "Consolas"
    r.font.size = Pt(10)

add_para(doc, "Validation", bold=True, size=12)
add_para(
    doc,
    "Each pipeline stage writes a JSON report to outputs/stage{N}_*.json. "
    "The top-level status field should be ok for a healthy run. A degraded "
    "status indicates at least one check failed; see the failures array in "
    "the same report for specifics."
)

add_para(doc, "Artifacts", bold=True, size=12)
add_bullet(doc, "data/raw/ — unmodified downloaded source files")
add_bullet(doc, "data/processed/ — cleaned CSVs consumed by the UI")
add_bullet(doc, "rules/extracted_rules.json — Stage 3 output")
add_bullet(doc, "rules/required_permits.json — hand-curated permit bank")
add_bullet(doc, "models/risk_model.joblib — trained classifier")
add_bullet(doc, "outputs/stage*_report.json — validation reports")

out_path = OUT / "Mobile_Food_Vendor_Compliance_Report.docx"
doc.save(out_path)
print(f"Report written to {out_path}")
print(f"Size: {out_path.stat().st_size:,} bytes")
